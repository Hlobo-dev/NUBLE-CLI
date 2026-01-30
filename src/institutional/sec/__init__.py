"""
SEC Filings Module - TENK Integration
======================================

AI-powered SEC filings analysis integrated from TENK project.
Uses Claude Opus 4.5 for analysis and local vector database for RAG.

Features:
- RAG over SEC filings (10-K, 10-Q, 8-K, 13F)
- Auto-downloads from SEC EDGAR
- Semantic search with sentence-transformers
- Local DuckDB vector database
- Export to PDF, DOCX, Excel
- Citation support with source links

Original TENK by Rallies.ai - Enhanced for Institutional Research Platform
"""

import os
import re
import json
import asyncio
from typing import Dict, List, Optional, Any, Tuple
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path

# Third-party imports (will gracefully degrade if not available)
try:
    import duckdb
    DUCKDB_AVAILABLE = True
except ImportError:
    DUCKDB_AVAILABLE = False

try:
    import numpy as np
    NUMPY_AVAILABLE = True
except ImportError:
    NUMPY_AVAILABLE = False

try:
    from sentence_transformers import SentenceTransformer
    SENTENCE_TRANSFORMERS_AVAILABLE = True
except ImportError:
    SENTENCE_TRANSFORMERS_AVAILABLE = False

try:
    from edgar import Company, set_identity
    EDGAR_AVAILABLE = True
except ImportError:
    EDGAR_AVAILABLE = False

try:
    import yfinance as yf
    YFINANCE_AVAILABLE = True
except ImportError:
    YFINANCE_AVAILABLE = False


# =============================================================================
# Configuration
# =============================================================================

@dataclass
class SECConfig:
    """Configuration for SEC filings module"""
    # Database
    db_path: str = "~/.institutional/sec_filings.db"
    
    # Search settings
    chunk_size: int = 3000
    chunk_overlap: int = 500
    top_k: int = 10
    
    # Embeddings
    embedding_model: str = "all-MiniLM-L6-v2"
    
    # EDGAR identity (required by SEC)
    edgar_identity: str = "InstitutionalResearch research@example.com"
    
    # Export directory
    exports_dir: str = "~/.institutional/exports"
    
    def __post_init__(self):
        self.db_path = os.path.expanduser(self.db_path)
        self.exports_dir = os.path.expanduser(self.exports_dir)
        os.makedirs(os.path.dirname(self.db_path), exist_ok=True)
        os.makedirs(self.exports_dir, exist_ok=True)


# =============================================================================
# Database Layer
# =============================================================================

class SECDatabase:
    """
    Local vector database for SEC filings using DuckDB.
    Stores chunked text with embeddings for semantic search.
    """
    
    def __init__(self, config: SECConfig = None):
        if not DUCKDB_AVAILABLE:
            raise ImportError("duckdb is required. Install with: pip install duckdb")
        if not SENTENCE_TRANSFORMERS_AVAILABLE:
            raise ImportError("sentence-transformers is required. Install with: pip install sentence-transformers")
        
        self.config = config or SECConfig()
        self.conn = duckdb.connect(self.config.db_path)
        self.model = SentenceTransformer(self.config.embedding_model)
        self._init_db()
    
    def _init_db(self):
        """Initialize database schema"""
        self.conn.execute("""
            CREATE TABLE IF NOT EXISTS filings (
                ticker VARCHAR,
                form VARCHAR,
                year INTEGER,
                quarter INTEGER,
                chunk_index INTEGER,
                chunk_text TEXT,
                embedding FLOAT[384],
                url VARCHAR,
                filing_date VARCHAR,
                PRIMARY KEY (ticker, form, year, quarter, chunk_index)
            )
        """)
        self.conn.execute("""
            CREATE INDEX IF NOT EXISTS idx_filing 
            ON filings (ticker, form, year, quarter)
        """)
    
    def chunk_text(self, text: str) -> List[str]:
        """Split text into overlapping chunks"""
        chunk_size = self.config.chunk_size
        overlap = self.config.chunk_overlap
        chunks = []
        start = 0
        while start < len(text):
            end = start + chunk_size
            chunks.append(text[start:end])
            start = end - overlap
        return [c for c in chunks if c.strip()]
    
    def has_filing(self, ticker: str, form: str, year: int, quarter: int = 0) -> bool:
        """Check if filing exists in database"""
        result = self.conn.execute("""
            SELECT 1 FROM filings
            WHERE ticker = ? AND form = ? AND year = ? AND quarter = ?
            LIMIT 1
        """, [ticker.upper(), form, year, quarter]).fetchone()
        return result is not None
    
    def add_filing(
        self,
        ticker: str,
        form: str,
        year: int,
        quarter: int,
        text: str,
        url: str = None,
        filing_date: str = None
    ):
        """Add a filing to the database"""
        ticker = ticker.upper()
        if self.has_filing(ticker, form, year, quarter):
            return
        
        chunks = self.chunk_text(text)
        embeddings = self.model.encode(chunks, show_progress_bar=False)
        
        self.conn.executemany(
            "INSERT INTO filings VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)",
            [
                (ticker, form, year, quarter, i, chunk, emb.tolist(), url, filing_date)
                for i, (chunk, emb) in enumerate(zip(chunks, embeddings))
            ]
        )
        return len(chunks)
    
    def search(
        self,
        query: str,
        k: int = None,
        ticker: str = None,
        form: str = None,
        year: int = None,
        quarter: int = None
    ) -> List[Dict[str, Any]]:
        """
        Semantic search over filings.
        
        Args:
            query: Natural language query
            k: Number of results (default: config.top_k)
            ticker: Filter by ticker
            form: Filter by form type (10-K, 10-Q, etc.)
            year: Filter by year
            quarter: Filter by quarter (0 for annual)
            
        Returns:
            List of matching chunks with metadata and scores
        """
        k = k or self.config.top_k
        
        # Build WHERE clause
        where = []
        params = []
        
        if ticker:
            where.append("ticker = ?")
            params.append(ticker.upper())
        if form:
            where.append("form = ?")
            params.append(form)
        if year:
            where.append("year = ?")
            params.append(year)
        if quarter is not None:
            where.append("quarter = ?")
            params.append(quarter)
        
        where_clause = f"WHERE {' AND '.join(where)}" if where else ""
        
        rows = self.conn.execute(f"""
            SELECT chunk_text, ticker, form, year, quarter, embedding, url, filing_date
            FROM filings {where_clause}
        """, params).fetchall()
        
        if not rows:
            return []
        
        texts = [r[0] for r in rows]
        meta = [
            {
                "ticker": r[1],
                "form": r[2],
                "year": r[3],
                "quarter": r[4],
                "url": r[6],
                "filing_date": r[7]
            }
            for r in rows
        ]
        embeddings = np.array([r[5] for r in rows])
        
        # Compute similarity
        query_emb = self.model.encode(query)
        sims = np.dot(embeddings, query_emb) / (
            np.linalg.norm(embeddings, axis=1) * np.linalg.norm(query_emb) + 1e-8
        )
        
        # Get top-k
        top_k = np.argsort(sims)[-k:][::-1]
        return [
            {"text": texts[i], **meta[i], "score": float(sims[i])}
            for i in top_k
        ]
    
    def list_filings(self, ticker: str = None) -> List[Dict[str, Any]]:
        """List all filings in database"""
        if ticker:
            rows = self.conn.execute("""
                SELECT ticker, form, year, quarter, COUNT(*) as chunks
                FROM filings
                WHERE ticker = ?
                GROUP BY ticker, form, year, quarter
                ORDER BY year DESC, quarter DESC
            """, [ticker.upper()]).fetchall()
        else:
            rows = self.conn.execute("""
                SELECT ticker, form, year, quarter, COUNT(*) as chunks
                FROM filings
                GROUP BY ticker, form, year, quarter
                ORDER BY ticker, year DESC, quarter DESC
            """).fetchall()
        
        return [
            {
                "ticker": r[0],
                "form": r[1],
                "year": r[2],
                "quarter": r[3] if r[3] != 0 else None,
                "chunks": r[4]
            }
            for r in rows
        ]
    
    def delete_filing(self, ticker: str, form: str, year: int, quarter: int = 0):
        """Delete a filing from database"""
        self.conn.execute("""
            DELETE FROM filings
            WHERE ticker = ? AND form = ? AND year = ? AND quarter = ?
        """, [ticker.upper(), form, year, quarter])
    
    def get_stats(self) -> Dict[str, Any]:
        """Get database statistics"""
        total_filings = self.conn.execute("""
            SELECT COUNT(DISTINCT (ticker, form, year, quarter)) FROM filings
        """).fetchone()[0]
        
        total_chunks = self.conn.execute("SELECT COUNT(*) FROM filings").fetchone()[0]
        
        tickers = self.conn.execute("""
            SELECT DISTINCT ticker FROM filings ORDER BY ticker
        """).fetchall()
        
        return {
            "total_filings": total_filings,
            "total_chunks": total_chunks,
            "tickers": [t[0] for t in tickers],
            "db_path": self.config.db_path
        }


# =============================================================================
# EDGAR Integration
# =============================================================================

class EDGARClient:
    """
    SEC EDGAR integration for fetching filings.
    """
    
    def __init__(self, config: SECConfig = None):
        if not EDGAR_AVAILABLE:
            raise ImportError("edgartools is required. Install with: pip install edgartools")
        
        self.config = config or SECConfig()
        set_identity(self.config.edgar_identity)
    
    def check_available(self, tickers: List[str]) -> Dict[str, List[Dict]]:
        """
        Check what filings are available on SEC EDGAR.
        
        Args:
            tickers: List of stock tickers
            
        Returns:
            Dict mapping ticker to list of available filings
        """
        results = {}
        
        for ticker in tickers:
            ticker = ticker.upper()
            company = Company(ticker)
            results[ticker] = []
            
            for form in ["10-K", "10-Q", "8-K"]:
                try:
                    filings = list(company.get_filings(form=form))[:10]
                    for f in filings:
                        report_date = getattr(f, 'report_date', None)
                        if not report_date:
                            continue
                        
                        parts = str(report_date).split("-")
                        if len(parts) < 2:
                            continue
                        
                        year = int(parts[0])
                        month = int(parts[1])
                        q = 1 if month <= 3 else 2 if month <= 6 else 3 if month <= 9 else 4
                        
                        results[ticker].append({
                            "form": form,
                            "year": year,
                            "quarter": q if form == "10-Q" else None,
                            "report_date": str(report_date),
                            "accession_number": getattr(f, 'accession_number', None)
                        })
                except Exception as e:
                    pass
        
        return results
    
    def download_filing(
        self,
        ticker: str,
        form: str,
        year: int,
        quarter: int = None
    ) -> Tuple[Optional[str], Optional[str], Optional[str]]:
        """
        Download a specific filing from EDGAR.
        
        Args:
            ticker: Stock ticker
            form: Filing form (10-K, 10-Q, 8-K)
            year: Filing year
            quarter: Quarter for 10-Q (1-4)
            
        Returns:
            Tuple of (text, url, filing_date) or (None, None, None) if not found
        """
        ticker = ticker.upper()
        company = Company(ticker)
        filings = list(company.get_filings(form=form))
        
        for f in filings:
            report_date = getattr(f, 'report_date', None)
            if not report_date:
                continue
            
            parts = str(report_date).split("-")
            if len(parts) < 2:
                continue
            
            f_year = int(parts[0])
            f_month = int(parts[1])
            f_quarter = 1 if f_month <= 3 else 2 if f_month <= 6 else 3 if f_month <= 9 else 4
            
            if form == "10-K" and f_year == year:
                text = f.text()
                if text and len(text.strip()) > 100:
                    return text, getattr(f, 'url', None), str(report_date)
            
            if form == "10-Q" and f_year == year and f_quarter == quarter:
                text = f.text()
                if text and len(text.strip()) > 100:
                    return text, getattr(f, 'url', None), str(report_date)
            
            if form == "8-K" and f_year == year:
                text = f.text()
                if text and len(text.strip()) > 100:
                    return text, getattr(f, 'url', None), str(report_date)
        
        return None, None, None
    
    def get_company_info(self, ticker: str) -> Dict[str, Any]:
        """Get company information from EDGAR"""
        ticker = ticker.upper()
        try:
            company = Company(ticker)
            return {
                "ticker": ticker,
                "name": getattr(company, 'name', ticker),
                "cik": getattr(company, 'cik', None),
                "sic": getattr(company, 'sic', None),
                "industry": getattr(company, 'industry', None),
            }
        except:
            return {"ticker": ticker, "name": ticker}


# =============================================================================
# SEC Filing Analyst (Claude-Powered)
# =============================================================================

@dataclass
class SECSearchResult:
    """Result from SEC filing search"""
    text: str
    ticker: str
    form: str
    year: int
    quarter: Optional[int]
    score: float
    url: Optional[str] = None
    filing_date: Optional[str] = None


@dataclass
class SECAnalysisResult:
    """Result from SEC analysis"""
    query: str
    answer: str
    sources: List[SECSearchResult]
    tickers: List[str]
    filings_used: List[Dict[str, Any]]
    tokens_used: int = 0
    latency_ms: float = 0


class SECFilingAnalyst:
    """
    AI-powered SEC filing analyst using Claude Opus 4.5.
    
    Combines:
    - EDGAR integration for fetching filings
    - Local vector database for semantic search (RAG)
    - Claude Opus 4.5 for intelligent analysis
    
    Features:
    - Ask questions about 10-K, 10-Q, 8-K filings
    - Compare multiple companies
    - Track specific metrics across filings
    - Generate insights with citations
    """
    
    SYSTEM_PROMPT = """You are an elite investment analyst at a top hedge fund specializing in SEC filings analysis. 
Your expertise includes 10-K, 10-Q, 8-K reports and you can find insights most analysts miss.

Today is {date}. Current year is {year}.

## Your Capabilities
- Analyze SEC filings with deep financial expertise
- Find specific data points in complex documents
- Compare companies across filings
- Identify risks, opportunities, and red flags
- Track financial metrics and trends

## Guidelines
1. Base answers on the provided filing excerpts
2. Include specific numbers and data points
3. Cite sources with [Source: TICKER Form Year Q#] format
4. If information isn't in the excerpts, say so clearly
5. Be concise but comprehensive (200-500 words unless more detail needed)
6. For financial tables, format clearly with proper alignment
7. Identify risks and opportunities when relevant
8. Compare to industry benchmarks when possible

## Important
- Always prioritize accuracy over speculation
- Acknowledge uncertainty when appropriate
- Provide actionable insights for investment decisions
"""

    def __init__(
        self,
        config: SECConfig = None,
        claude_synthesizer = None
    ):
        """
        Initialize SEC Filing Analyst.
        
        Args:
            config: SECConfig for database and search settings
            claude_synthesizer: ClaudeSynthesizer instance (will create if not provided)
        """
        self.config = config or SECConfig()
        
        # Initialize database
        self.db = SECDatabase(self.config)
        
        # Initialize EDGAR client
        if EDGAR_AVAILABLE:
            self.edgar = EDGARClient(self.config)
        else:
            self.edgar = None
        
        # Initialize Claude synthesizer
        if claude_synthesizer:
            self.claude = claude_synthesizer
        else:
            try:
                from institutional.core.claude_synthesizer import create_claude_synthesizer
                self.claude = create_claude_synthesizer(model="opus")
            except:
                self.claude = None
    
    def load_filing(
        self,
        ticker: str,
        form: str,
        year: int,
        quarter: int = None
    ) -> Dict[str, Any]:
        """
        Load a filing into the local database.
        
        Args:
            ticker: Stock ticker
            form: 10-K, 10-Q, or 8-K
            year: Filing year
            quarter: Quarter for 10-Q (1-4)
            
        Returns:
            Status dict with loading result
        """
        ticker = ticker.upper()
        q = quarter if form == "10-Q" else 0
        
        # Check if already loaded
        if self.db.has_filing(ticker, form, year, q):
            return {
                "status": "already_loaded",
                "ticker": ticker,
                "form": form,
                "year": year,
                "quarter": quarter
            }
        
        if not self.edgar:
            return {
                "status": "error",
                "error": "EDGAR client not available. Install edgartools."
            }
        
        # Download from EDGAR
        text, url, filing_date = self.edgar.download_filing(ticker, form, year, quarter)
        
        if not text:
            return {
                "status": "not_found",
                "ticker": ticker,
                "form": form,
                "year": year,
                "quarter": quarter
            }
        
        # Add to database
        chunks = self.db.add_filing(ticker, form, year, q, text, url, filing_date)
        
        return {
            "status": "loaded",
            "ticker": ticker,
            "form": form,
            "year": year,
            "quarter": quarter,
            "chunks": chunks,
            "url": url
        }
    
    def search(
        self,
        queries: List[str],
        ticker: str = None,
        year: int = None,
        quarter: int = None
    ) -> Dict[str, List[SECSearchResult]]:
        """
        Semantic search over loaded filings.
        
        Args:
            queries: List of search queries (use multiple phrasings for better results)
            ticker: Filter by ticker
            year: Filter by year
            quarter: Filter by quarter
            
        Returns:
            Dict mapping query to list of results
        """
        results = {}
        
        for query in queries:
            db_results = self.db.search(
                query=query,
                ticker=ticker,
                year=year,
                quarter=quarter
            )
            
            results[query] = [
                SECSearchResult(
                    text=r["text"],
                    ticker=r["ticker"],
                    form=r["form"],
                    year=r["year"],
                    quarter=r["quarter"] if r["quarter"] != 0 else None,
                    score=r["score"],
                    url=r.get("url"),
                    filing_date=r.get("filing_date")
                )
                for r in db_results
            ]
        
        return results
    
    def ask(
        self,
        question: str,
        tickers: List[str] = None,
        year: int = None,
        auto_load: bool = True
    ) -> SECAnalysisResult:
        """
        Ask a question about SEC filings.
        
        Args:
            question: Natural language question
            tickers: List of tickers to search (extracts from question if not provided)
            year: Filter to specific year
            auto_load: Automatically load latest filings if not in database
            
        Returns:
            SECAnalysisResult with answer and sources
        """
        import time
        start = time.time()
        
        # Extract tickers from question if not provided
        if not tickers:
            tickers = self._extract_tickers(question)
        
        if not tickers:
            return SECAnalysisResult(
                query=question,
                answer="Please specify which company/ticker you want to analyze.",
                sources=[],
                tickers=[],
                filings_used=[]
            )
        
        # Auto-load latest filings if needed
        if auto_load:
            current_year = datetime.now().year
            for ticker in tickers:
                # Try to load latest 10-K
                if not self.db.has_filing(ticker, "10-K", current_year, 0):
                    self.load_filing(ticker, "10-K", current_year)
                if not self.db.has_filing(ticker, "10-K", current_year - 1, 0):
                    self.load_filing(ticker, "10-K", current_year - 1)
        
        # Generate search queries
        search_queries = self._generate_search_queries(question)
        
        # Search filings
        all_results = []
        for ticker in tickers:
            search_results = self.search(
                queries=search_queries,
                ticker=ticker,
                year=year
            )
            for query_results in search_results.values():
                all_results.extend(query_results)
        
        # Deduplicate and sort by score
        seen = set()
        unique_results = []
        for r in sorted(all_results, key=lambda x: x.score, reverse=True):
            key = (r.ticker, r.form, r.year, r.quarter, r.text[:100])
            if key not in seen:
                seen.add(key)
                unique_results.append(r)
        
        # Take top results
        top_results = unique_results[:15]
        
        if not top_results:
            return SECAnalysisResult(
                query=question,
                answer=f"No relevant information found in loaded filings for {', '.join(tickers)}. Try loading more filings with load_filing().",
                sources=[],
                tickers=tickers,
                filings_used=[]
            )
        
        # Build context for Claude
        context = self._build_context(question, top_results)
        
        # Get Claude's analysis
        if self.claude:
            now = datetime.now()
            system = self.SYSTEM_PROMPT.format(
                date=now.strftime("%B %d, %Y"),
                year=now.year
            )
            
            response = self.claude._call_claude(context, system_prompt=system)
            answer = response.content
            tokens = response.input_tokens + response.output_tokens
        else:
            # Fallback: just return the raw excerpts
            answer = "Claude not available. Here are the relevant excerpts:\n\n"
            for i, r in enumerate(top_results[:5], 1):
                answer += f"**[{r.ticker} {r.form} {r.year}]**\n{r.text[:500]}...\n\n"
            tokens = 0
        
        latency = (time.time() - start) * 1000
        
        # Get unique filings used
        filings_used = []
        seen_filings = set()
        for r in top_results:
            key = (r.ticker, r.form, r.year, r.quarter)
            if key not in seen_filings:
                seen_filings.add(key)
                filings_used.append({
                    "ticker": r.ticker,
                    "form": r.form,
                    "year": r.year,
                    "quarter": r.quarter,
                    "url": r.url
                })
        
        return SECAnalysisResult(
            query=question,
            answer=answer,
            sources=top_results,
            tickers=tickers,
            filings_used=filings_used,
            tokens_used=tokens,
            latency_ms=latency
        )
    
    def compare(
        self,
        tickers: List[str],
        aspect: str,
        year: int = None
    ) -> SECAnalysisResult:
        """
        Compare multiple companies on a specific aspect.
        
        Args:
            tickers: List of tickers to compare
            aspect: What to compare (e.g., "revenue growth", "risk factors", "margins")
            year: Year to compare (defaults to latest)
            
        Returns:
            SECAnalysisResult with comparison
        """
        question = f"Compare {', '.join(tickers)} on {aspect}. Provide specific numbers and data points."
        return self.ask(question, tickers=tickers, year=year)
    
    def get_risk_factors(self, ticker: str, year: int = None) -> SECAnalysisResult:
        """Get risk factors from 10-K filing"""
        question = f"What are the main risk factors disclosed by {ticker}? Summarize Item 1A."
        return self.ask(question, tickers=[ticker], year=year)
    
    def get_business_overview(self, ticker: str, year: int = None) -> SECAnalysisResult:
        """Get business overview from 10-K filing"""
        question = f"Describe {ticker}'s business model, products/services, and competitive position. Summarize Item 1."
        return self.ask(question, tickers=[ticker], year=year)
    
    def get_mda(self, ticker: str, year: int = None) -> SECAnalysisResult:
        """Get Management Discussion & Analysis from 10-K"""
        question = f"Summarize {ticker}'s MD&A section (Item 7). Focus on key financial trends, management outlook, and notable changes."
        return self.ask(question, tickers=[ticker], year=year)
    
    def _extract_tickers(self, text: str) -> List[str]:
        """Extract stock tickers from text"""
        # Common patterns for tickers
        pattern = r'\b([A-Z]{1,5})\b'
        matches = re.findall(pattern, text)
        
        # Filter out common words
        common_words = {'A', 'I', 'IT', 'AT', 'TO', 'IN', 'ON', 'OR', 'AN', 'AS', 'BY', 'IS', 'IF', 'BE', 'WE', 'DO', 'GO', 'NO', 'SO', 'UP', 'FOR', 'THE', 'AND', 'BUT', 'NOT', 'YOU', 'ALL', 'CAN', 'HER', 'WAS', 'ONE', 'OUR', 'OUT', 'SEC', 'CEO', 'CFO', 'COO', 'USA', 'GDP', 'IPO', 'ETF', 'NYSE', 'NASDAQ', 'Q1', 'Q2', 'Q3', 'Q4', 'YOY', 'QOQ', 'EBITDA', 'EPS', 'PE', 'PS'}
        
        tickers = [m for m in matches if m not in common_words and len(m) >= 2]
        return list(dict.fromkeys(tickers))  # Remove duplicates, preserve order
    
    def _generate_search_queries(self, question: str) -> List[str]:
        """Generate multiple search queries from a question"""
        queries = [question]
        
        # Add variations based on common financial terms
        q_lower = question.lower()
        
        if 'revenue' in q_lower or 'sales' in q_lower:
            queries.extend(['net revenue', 'total sales', 'revenue growth'])
        if 'profit' in q_lower or 'margin' in q_lower:
            queries.extend(['gross profit', 'operating margin', 'net income margin'])
        if 'risk' in q_lower:
            queries.extend(['Item 1A Risk Factors', 'material risks', 'risk factors'])
        if 'competition' in q_lower or 'competitor' in q_lower:
            queries.extend(['competitive landscape', 'key competitors', 'market share'])
        if 'business' in q_lower or 'company' in q_lower:
            queries.extend(['Item 1 Business', 'business description', 'our business'])
        if 'guidance' in q_lower or 'outlook' in q_lower:
            queries.extend(['forward-looking', 'management outlook', 'expect', 'anticipate'])
        
        return queries[:5]  # Limit to 5 queries
    
    def _build_context(self, question: str, results: List[SECSearchResult]) -> str:
        """Build context message for Claude"""
        context = f"# Question\n{question}\n\n"
        context += "# Relevant Filing Excerpts\n\n"
        
        for i, r in enumerate(results, 1):
            q_str = f" Q{r.quarter}" if r.quarter else ""
            context += f"## [{i}] {r.ticker} {r.form} {r.year}{q_str}\n"
            if r.url:
                context += f"Source: {r.url}\n"
            context += f"Relevance Score: {r.score:.3f}\n\n"
            context += f"```\n{r.text}\n```\n\n"
        
        context += "---\n\nBased on the above excerpts, please answer the question. Cite sources using [Source: TICKER Form Year] format."
        
        return context
    
    # Convenience methods
    def list_loaded(self, ticker: str = None) -> List[Dict]:
        """List loaded filings"""
        return self.db.list_filings(ticker)
    
    def check_available(self, tickers: List[str]) -> Dict[str, List[Dict]]:
        """Check available filings on EDGAR"""
        if self.edgar:
            return self.edgar.check_available(tickers)
        return {}
    
    def get_stats(self) -> Dict[str, Any]:
        """Get database statistics"""
        return self.db.get_stats()


# =============================================================================
# Export Functions
# =============================================================================

def export_to_pdf(content: str, filepath: str = None) -> str:
    """Export markdown content to PDF"""
    try:
        from markdown import markdown
        from weasyprint import HTML
        
        if not filepath:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filepath = f"sec_analysis_{timestamp}.pdf"
        
        html_content = markdown(content, extensions=['tables', 'fenced_code'])
        html = f"""
        <html>
        <head>
            <style>
                body {{ font-family: 'Times New Roman', Times, serif; font-size: 12pt; }}
                table {{ border-collapse: collapse; width: 100%; margin: 10pt 0; }}
                th, td {{ border: 1px solid #000; padding: 4pt 6pt; }}
                pre {{ background: #f5f5f5; padding: 8pt; font-size: 10pt; }}
            </style>
        </head>
        <body>{html_content}</body>
        </html>
        """
        HTML(string=html).write_pdf(filepath)
        return filepath
    except ImportError:
        raise ImportError("weasyprint and markdown required for PDF export")


def export_to_docx(content: str, filepath: str = None) -> str:
    """Export markdown content to DOCX"""
    try:
        from docx import Document
        from docx.shared import Pt
        
        if not filepath:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filepath = f"sec_analysis_{timestamp}.docx"
        
        doc = Document()
        
        # Simple markdown parsing
        for line in content.split('\n'):
            if line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.strip():
                doc.add_paragraph(line)
        
        doc.save(filepath)
        return filepath
    except ImportError:
        raise ImportError("python-docx required for DOCX export")


# =============================================================================
# Convenience Functions
# =============================================================================

def create_sec_analyst(
    config: SECConfig = None,
    claude_model: str = "opus"
) -> SECFilingAnalyst:
    """
    Create an SEC Filing Analyst instance.
    
    Args:
        config: SECConfig (uses defaults if not provided)
        claude_model: "opus", "sonnet", or "haiku"
        
    Returns:
        SECFilingAnalyst instance
    """
    try:
        from institutional.core.claude_synthesizer import create_claude_synthesizer
        claude = create_claude_synthesizer(model=claude_model)
    except:
        claude = None
    
    return SECFilingAnalyst(config=config, claude_synthesizer=claude)


# Quick access function
def ask_sec(question: str, tickers: List[str] = None) -> str:
    """
    Quick function to ask a question about SEC filings.
    
    Args:
        question: Your question
        tickers: Optional list of tickers (will extract from question if not provided)
        
    Returns:
        Answer string
    """
    analyst = create_sec_analyst()
    result = analyst.ask(question, tickers=tickers)
    return result.answer
